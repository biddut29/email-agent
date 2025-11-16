"""
Text Extraction Utility - Extract text from images (OCR) and PDFs
"""

import io
from typing import Optional, Dict
import logging

logger = logging.getLogger(__name__)


def extract_text_from_image(binary_data: bytes, filename: str = None) -> Optional[str]:
    """
    Extract text from image using OCR (Optical Character Recognition)
    
    Args:
        binary_data: Image binary data
        filename: Optional filename for logging
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        from PIL import Image
        import pytesseract
        
        # Load image from binary data
        image = Image.open(io.BytesIO(binary_data))
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        
        # Clean up text (remove extra whitespace)
        text = text.strip()
        
        if text:
            print(f"✅ OCR extracted {len(text)} characters from image: {filename or 'unknown'}")
            return text
        else:
            print(f"⚠️  OCR returned no text from image: {filename or 'unknown'}")
            return None
            
    except ImportError as e:
        print(f"⚠️  OCR dependencies not installed: {e}")
        print("   Install Tesseract OCR: https://github.com/tesseract-ocr/tesseract")
        print("   Then install: pip install pytesseract Pillow")
        return None
    except Exception as e:
        print(f"❌ OCR extraction failed for {filename or 'unknown'}: {e}")
        return None


def extract_text_from_pdf(binary_data: bytes, filename: str = None) -> Optional[str]:
    """
    Extract text from PDF file
    
    Args:
        binary_data: PDF binary data
        filename: Optional filename for logging
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(binary_data)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    text = "\n\n".join(text_parts)
                    text = text.strip()
                    if text:
                        print(f"✅ PDF extracted {len(text)} characters from: {filename or 'unknown'}")
                        return text
        except ImportError:
            # Fallback to PyPDF2
            pass
        except Exception as e:
            print(f"⚠️  pdfplumber extraction failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(binary_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    print(f"⚠️  Error extracting text from PDF page {page_num + 1}: {e}")
                    continue
            
            if text_parts:
                text = "\n\n".join(text_parts)
                text = text.strip()
                if text:
                    print(f"✅ PDF extracted {len(text)} characters (PyPDF2) from: {filename or 'unknown'}")
                    return text
        except ImportError:
            print(f"⚠️  PDF extraction libraries not installed")
            print("   Install: pip install pdfplumber PyPDF2")
            return None
        except Exception as e:
            print(f"❌ PDF extraction failed for {filename or 'unknown'}: {e}")
            return None
        
        print(f"⚠️  No text extracted from PDF: {filename or 'unknown'}")
        return None
        
    except Exception as e:
        print(f"❌ PDF extraction error for {filename or 'unknown'}: {e}")
        return None


def extract_text_from_word(binary_data: bytes, filename: str = None) -> Optional[str]:
    """
    Extract text from Word document (.docx)
    
    Args:
        binary_data: Word document binary data
        filename: Optional filename for logging
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(binary_data))
        
        text_parts = []
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if text_parts:
            text = "\n".join(text_parts)
            text = text.strip()
            if text:
                print(f"✅ Word document extracted {len(text)} characters from: {filename or 'unknown'}")
                return text
        
        print(f"⚠️  No text extracted from Word document: {filename or 'unknown'}")
        return None
        
    except ImportError:
        print(f"⚠️  python-docx not installed")
        print("   Install: pip install python-docx")
        return None
    except Exception as e:
        print(f"❌ Word document extraction failed for {filename or 'unknown'}: {e}")
        return None


def extract_text_from_excel(binary_data: bytes, filename: str = None) -> Optional[str]:
    """
    Extract text from Excel file (.xlsx)
    
    Args:
        binary_data: Excel file binary data
        filename: Optional filename for logging
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        from openpyxl import load_workbook
        
        workbook = load_workbook(io.BytesIO(binary_data), data_only=True)
        
        text_parts = []
        
        # Extract text from all sheets
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = [f"Sheet: {sheet_name}"]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell_value in row:
                    if cell_value is not None:
                        # Convert cell value to string
                        cell_str = str(cell_value).strip()
                        if cell_str:
                            row_text.append(cell_str)
                if row_text:
                    sheet_text.append(" | ".join(row_text))
            
            if len(sheet_text) > 1:  # More than just the sheet name
                text_parts.append("\n".join(sheet_text))
        
        if text_parts:
            text = "\n\n".join(text_parts)
            text = text.strip()
            if text:
                print(f"✅ Excel file extracted {len(text)} characters from: {filename or 'unknown'}")
                return text
        
        print(f"⚠️  No text extracted from Excel file: {filename or 'unknown'}")
        return None
        
    except ImportError:
        print(f"⚠️  openpyxl not installed")
        print("   Install: pip install openpyxl")
        return None
    except Exception as e:
        print(f"❌ Excel extraction failed for {filename or 'unknown'}: {e}")
        return None


def extract_text_from_attachment(attachment: Dict, binary_data: bytes) -> Optional[str]:
    """
    Extract text from attachment based on content type
    
    Args:
        attachment: Attachment dictionary with filename, content_type, etc.
        binary_data: Attachment binary data
        
    Returns:
        Extracted text or None
    """
    if not binary_data:
        return None
    
    filename = attachment.get('filename', '').lower()
    content_type = attachment.get('content_type', '').lower()
    
    # Check if it's an image (OCR)
    if content_type.startswith('image/') or any(filename.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']):
        return extract_text_from_image(binary_data, attachment.get('filename', ''))
    
    # Check if it's a PDF
    elif content_type == 'application/pdf' or filename.endswith('.pdf'):
        return extract_text_from_pdf(binary_data, attachment.get('filename', ''))
    
    # Check if it's a Word document (.docx)
    elif (content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
          filename.endswith('.docx')):
        return extract_text_from_word(binary_data, attachment.get('filename', ''))
    
    # Check if it's an Excel file (.xlsx)
    elif (content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' or 
          filename.endswith('.xlsx')):
        return extract_text_from_excel(binary_data, attachment.get('filename', ''))
    
    # For text files, text_content should already be set during email parsing
    # But we can also try to extract if binary_data is available
    elif content_type.startswith('text/'):
        try:
            text = binary_data.decode('utf-8', errors='ignore')
            if text.strip():
                return text.strip()
        except:
            pass
    
    return None

